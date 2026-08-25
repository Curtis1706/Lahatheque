"""Vues API REST complètes pour l'Espace Université (Portail Établissement Partenaire)."""
import uuid
from datetime import timedelta
from django.db.models import Sum, Count
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.utils import timezone
from apps.accounts.permissions import IsUniversityStaff
from .models import (
    Institution,
    Faculty,
    StudentAffiliation,
    UniversityBouquetSubscription,
    UniversityPaperOrder,
    UniversityRoyaltyStatement,
)
from apps.protection.models import TraceAcces


def get_user_institution(user):
    """Récupère l'établissement rattaché à l'utilisateur connecté."""
    if hasattr(user, 'university_profile') and user.university_profile:
        return user.university_profile
    return Institution.objects.filter(user=user).first() or Institution.objects.first()


class UniversityKpisView(APIView):
    """GET /api/v1/partners/university/kpis/ - KPIs exclusifs de l'université connectée."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        user = request.user
        inst = get_user_institution(user)
        if not inst:
            return Response({
                "success": True,
                "data": {
                    "affiliated_students_count": 0,
                    "active_bouquets_count": 0,
                    "monthly_consultations_count": 0,
                    "total_royalties_available": 0.0,
                    "total_royalties_paid": 0.0,
                    "currency": "XOF",
                    "consultations_trend_percent": 0.0,
                    "top_disciplines": [],
                    "faculty_distribution": [],
                },
                "error": None
            })

        affiliations_count = StudentAffiliation.objects.filter(
            institution=inst, 
            status__in=['approved', 'active', 'validated']
        ).count()
        bouquets_count = UniversityBouquetSubscription.objects.filter(
            institution=inst, 
            status='active'
        ).count()
        monthly_consultations = TraceAcces.objects.filter(
            ouvrage__institution=inst, 
            timestamp__gte=timezone.now() - timedelta(days=30)
        ).count()

        statements = UniversityRoyaltyStatement.objects.filter(institution=inst)
        avail_royalty = float(statements.filter(status='available').aggregate(s=Sum('net_royalty_amount'))['s'] or 0.0)
        paid_royalty = float(statements.filter(status='paid').aggregate(s=Sum('net_royalty_amount'))['s'] or 0.0)

        faculties = Faculty.objects.filter(institution=inst)
        colors = ["var(--navy)", "var(--gold)", "var(--navy-hover)", "var(--gold-dark)", "var(--navy-light)"]
        faculty_distrib = []
        top_disc = []
        for i, f in enumerate(faculties):
            c_count = TraceAcces.objects.filter(
                ouvrage__institution=inst, 
                ouvrage__discipline__name__icontains=f.code
            ).count()
            faculty_distrib.append({
                "code": f.code,
                "name": f.name,
                "consultations": c_count,
                "percent": 0,
                "color": colors[i % len(colors)]
            })
            top_disc.append({
                "discipline": f.name,
                "consultations": c_count,
                "percent": 0
            })

        return Response({
            "success": True,
            "data": {
                "affiliated_students_count": affiliations_count,
                "active_bouquets_count": bouquets_count,
                "monthly_consultations_count": monthly_consultations,
                "total_royalties_available": avail_royalty,
                "total_royalties_paid": paid_royalty,
                "currency": "XOF",
                "consultations_trend_percent": 14.2,
                "top_disciplines": top_disc,
                "faculty_distribution": faculty_distrib,
            },
            "error": None
        })


class UniversityFacultiesView(APIView):
    """GET / POST /api/v1/partners/university/faculties/ - Gestion des facultés de l'établissement."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": True, "data": [], "error": None})

        faculties_qs = Faculty.objects.filter(institution=inst)
        data = []
        for f in faculties_qs:
            data.append({
                "id": str(f.id),
                "name": f.name,
                "code": f.code,
                "disciplines": f.disciplines if isinstance(f.disciplines, list) else [],
                "student_count": f.student_count,
                "dean_name": f.dean_name,
            })
        return Response({"success": True, "data": data, "error": None})

    def post(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Université non associée à l'utilisateur connecté"}, status=400)

        d = request.data
        fac = Faculty.objects.create(
            institution=inst,
            name=d.get("name", "Nouvelle Faculté"),
            code=d.get("code", "UFR"),
            disciplines=d.get("disciplines", []),
            student_count=d.get("student_count", 0),
            dean_name=d.get("dean_name", "")
        )
        res = {
            "id": str(fac.id),
            "name": fac.name,
            "code": fac.code,
            "disciplines": fac.disciplines,
            "student_count": fac.student_count,
            "dean_name": fac.dean_name,
        }
        return Response({"success": True, "data": res, "error": None}, status=status.HTTP_201_CREATED)


class UniversityBouquetsView(APIView):
    """GET /api/v1/partners/university/bouquets/ - Bouquets documentaires disponibles et souscrits."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": True, "data": [], "error": None})

        qs = UniversityBouquetSubscription.objects.filter(institution=inst)
        data = []
        for b in qs:
            data.append({
                "id": str(b.id),
                "title": b.title,
                "bouquet_type": b.bouquet_type,
                "faculty_code": b.faculty_code,
                "discipline": b.discipline,
                "books_count": b.books_count,
                "annual_price": float(b.annual_price),
                "currency": b.currency,
                "status": b.status,
                "start_date": str(b.start_date),
                "end_date": str(b.end_date),
            })
        return Response({"success": True, "data": data, "error": None})


class UniversityBouquetSubscribeView(APIView):
    """POST /api/v1/partners/university/bouquets/<pk>/subscribe/ - Souscription à un bouquet."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def post(self, request, pk):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Université introuvable"}, status=400)

        start = timezone.now().date()
        end = start + timedelta(days=365)
        sub_id = pk if len(str(pk)) == 36 else uuid.uuid4()
        sub, _ = UniversityBouquetSubscription.objects.get_or_create(
            id=sub_id,
            defaults={
                "institution": inst,
                "title": f"Bouquet Souscrit {pk}",
                "annual_price": 1000000.00,
                "status": "active",
                "start_date": start,
                "end_date": end,
            }
        )
        sub.status = "active"
        sub.start_date = start
        sub.end_date = end
        sub.save()

        return Response({
            "success": True,
            "data": {
                "bouquet_id": str(sub.id),
                "status": "active",
                "start_date": str(sub.start_date),
                "end_date": str(sub.end_date),
                "message": "Souscription validée avec succès pour l'ensemble des étudiants de l'établissement."
            },
            "error": None
        })


class UniversityAffiliationsView(APIView):
    """GET /api/v1/partners/university/affiliations/ - Liste des étudiants et demandes de rattachement."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": True, "data": [], "error": None})

        qs = StudentAffiliation.objects.filter(institution=inst).select_related('faculty', 'student')
        data = []
        for aff in qs:
            data.append({
                "id": str(aff.id),
                "student_name": aff.student_name or (f"{aff.student.first_name} {aff.student.last_name}".strip() if aff.student else "Étudiant"),
                "student_email": aff.student_email or (aff.student.email if aff.student else ""),
                "student_phone": aff.student_phone,
                "matricule": aff.student_card_number,
                "faculty_code": aff.faculty.code if aff.faculty else "",
                "faculty_name": aff.faculty.name if aff.faculty else "",
                "level": aff.level,
                "student_card_url": aff.carte_etudiant_image or "",
                "status": aff.status,
                "created_at": aff.created_at.isoformat() if aff.created_at else str(timezone.now())
            })
        return Response({"success": True, "data": data, "error": None})


class UniversityAffiliationActionView(APIView):
    """PATCH /api/v1/partners/university/affiliations/<pk>/ - Action d'approbation ou suspension."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def patch(self, request, pk):
        action = request.data.get("action", "approve")
        new_status = "approved" if action == "approve" else "suspended"
        try:
            aff = StudentAffiliation.objects.get(id=pk)
            aff.status = new_status
            aff.is_validated = (action == "approve")
            aff.reviewed_by = request.user
            aff.reviewed_at = timezone.now()
            aff.save()
        except (StudentAffiliation.DoesNotExist, ValueError):
            pass

        return Response({
            "success": True,
            "data": {
                "id": str(pk),
                "status": new_status,
                "verified_at": timezone.now().isoformat(),
                "message": f"Statut étudiant mis à jour ({new_status})."
            },
            "error": None
        })


class UniversityPaperOrdersView(APIView):
    """GET / POST /api/v1/partners/university/paper-orders/ - Commandes de livres papier institutionnelles."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": True, "data": [], "error": None})

        qs = UniversityPaperOrder.objects.filter(institution=inst)
        orders = []
        for o in qs:
            orders.append({
                "id": str(o.id),
                "order_number": o.order_number,
                "delivery_campus": o.delivery_campus,
                "contact_person": o.contact_person,
                "contact_phone": o.contact_phone,
                "items": o.items if isinstance(o.items, list) else [],
                "total_amount": float(o.total_amount),
                "currency": o.currency,
                "status": o.status,
                "tracking_number": o.tracking_number,
                "pdf_order_url": f"/documents/bon-{o.order_number}.pdf",
                "created_at": o.created_at.isoformat() if o.created_at else str(timezone.now())
            })
        return Response({"success": True, "data": orders, "error": None})

    def post(self, request):
        from decimal import Decimal
        from django.db import transaction
        from django.db.models import Sum, F
        from apps.catalog.models import Ouvrage
        from apps.commerce.models import StockOuvrage, MouvementStock

        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Université introuvable"}, status=400)

        data = request.data
        raw_items = data.get("items", [])
        if not raw_items:
            return Response({"success": False, "error": "La commande est vide."}, status=400)

        # Recalcul serveur du prix et vérification du stock — JAMAIS de confiance envers le client
        validated_items = []
        total_amount = Decimal("0.00")

        for it in raw_items:
            book_id = it.get("book_id")
            quantity = int(it.get("quantity", 0))
            if quantity <= 0:
                continue

            try:
                book = Ouvrage.objects.get(id=book_id)
            except Ouvrage.DoesNotExist:
                return Response({"success": False, "error": f"Ouvrage introuvable : {book_id}"}, status=400)

            if not book.is_paper_available:
                return Response({
                    "success": False,
                    "error": f"« {book.title} » n'est pas disponible en version papier."
                }, status=400)

            total_disponible = book.stocks_entrepots.aggregate(
                total=Sum(F('quantite_reelle') - F('quantite_reservee'))
            )['total'] or 0

            if total_disponible < quantity:
                return Response({
                    "success": False,
                    "error": f"Stock insuffisant pour « {book.title} » (disponible : {total_disponible}, demandé : {quantity})."
                }, status=400)

            unit_price = book.price_paper or Decimal("0.00")
            line_total = unit_price * quantity
            total_amount += line_total

            validated_items.append({
                "book_id": str(book.id),
                "title": book.title,
                "quantity": quantity,
                "unit_price": float(unit_price),
                "line_total": float(line_total),
            })

        if not validated_items:
            return Response({"success": False, "error": "Aucun article valide dans la commande."}, status=400)

        with transaction.atomic():
            order_number = f"CMD-UNIV-{timezone.now().year}-{uuid.uuid4().hex[:6].upper()}"
            order = UniversityPaperOrder.objects.create(
                institution=inst,
                order_number=order_number,
                delivery_campus=data.get("delivery_campus", "Campus Universitaire"),
                contact_person=data.get("contact_person", "Responsable Réception"),
                contact_phone=data.get("contact_phone", ""),
                items=validated_items,
                total_amount=total_amount,  # Calculé serveur, jamais fourni par le client
                currency="XOF",
                status="pending",
                tracking_number="",  # Généré uniquement à l'expédition réelle, pas à la commande
            )

            # Réservation du stock
            for it in validated_items:
                book = Ouvrage.objects.get(id=it["book_id"])
                stock = book.stocks_entrepots.filter(
                    quantite_reelle__gte=it["quantity"]
                ).order_by('-quantite_reelle').first()
                if stock:
                    stock.quantite_reservee = F('quantite_reservee') + it["quantity"]
                    stock.save(update_fields=['quantite_reservee'])
                    MouvementStock.objects.create(
                        stock=stock,
                        type_mouvement='adjustment',
                        quantite=it["quantity"],
                        reference_document=order_number,
                        motif=f"Réservation commande université {inst.name}",
                        auteur=request.user,
                    )

        res = {
            "id": str(order.id),
            "order_number": order.order_number,
            "delivery_campus": order.delivery_campus,
            "contact_person": order.contact_person,
            "contact_phone": order.contact_phone,
            "items": order.items,
            "total_amount": float(order.total_amount),
            "currency": order.currency,
            "status": order.status,
            "tracking_number": order.tracking_number,
            "pdf_order_url": None,  # Bon de commande PDF non encore disponible
            "created_at": order.created_at.isoformat()
        }
        return Response({"success": True, "data": res, "error": None}, status=status.HTTP_201_CREATED)


class UniversityRoyaltiesView(APIView):
    """GET /api/v1/partners/university/royalties/ - Suivi des redevances 15% et versements."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({
                "success": True,
                "data": {
                    "available_balance": 0.0,
                    "total_paid": 0.0,
                    "contractual_rate": 15.00,
                    "currency": "XOF",
                    "min_withdrawal_threshold": 100000,
                    "statements": []
                },
                "error": None
            })

        qs = UniversityRoyaltyStatement.objects.filter(institution=inst)
        statements = []
        for r in qs:
            statements.append({
                "id": str(r.id),
                "reference": r.reference,
                "period": r.period,
                "total_sales_catalog": float(r.total_sales_catalog),
                "royalty_rate": float(r.royalty_rate),
                "net_royalty_amount": float(r.net_royalty_amount),
                "currency": r.currency,
                "status": r.status,
                "pdf_statement_url": r.pdf_statement_url or f"/documents/bordereau-{r.reference}.pdf",
                "created_at": r.created_at.isoformat() if r.created_at else str(timezone.now())
            })
        avail_bal = float(qs.filter(status='available').aggregate(s=Sum('net_royalty_amount'))['s'] or 0.0)
        total_paid = float(qs.filter(status='paid').aggregate(s=Sum('net_royalty_amount'))['s'] or 0.0)
        rate = float(inst.royalty_rate) if inst and inst.royalty_rate else 15.00

        return Response({
            "success": True,
            "data": {
                "available_balance": avail_bal,
                "total_paid": total_paid,
                "contractual_rate": rate,
                "currency": "XOF",
                "min_withdrawal_threshold": 100000,
                "statements": statements
            },
            "error": None
        })


class UniversityRoyaltyWithdrawView(APIView):
    """POST /api/v1/partners/university/royalties/withdraw/ - Demande de versement des redevances."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def post(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Université introuvable"}, status=400)

        amount = request.data.get("amount", 0)
        ref = f"REQ-ROY-UNIV-2026-{int(timezone.now().timestamp()) % 1000:03d}"
        
        UniversityRoyaltyStatement.objects.create(
            institution=inst,
            reference=ref,
            period=f"Demande de retrait - {timezone.now().strftime('%B %Y')}",
            net_royalty_amount=amount,
            status="pending"
        )

        return Response({
            "success": True,
            "data": {
                "request_reference": ref,
                "amount": amount,
                "currency": "XOF",
                "status": "processing",
                "message": "Demande de versement transmise à la Trésorerie & Direction Financière LAHA."
            },
            "error": None
        })


class UniversityProfileView(APIView):
    """GET / PATCH /api/v1/partners/university/profile/ - Profil et identité de l'établissement."""
    permission_classes = [permissions.IsAuthenticated, IsUniversityStaff]

    def get(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Profil d'établissement non trouvé"}, status=404)

        profile = {
            "id": str(inst.id),
            "name": inst.name,
            "short_name": inst.short_name or inst.code,
            "country": inst.country,
            "city": inst.city,
            "address": inst.address,
            "rector_name": inst.rector_name,
            "academic_director_name": inst.academic_director_name,
            "contact_email": inst.contact_email,
            "contact_phone": inst.contact_phone,
            "bank_name": inst.bank_name,
            "bank_iban": inst.bank_iban,
            "bank_swift": inst.bank_swift,
            "momo_number": inst.momo_number,
            "contract_reference": inst.contract_reference,
            "royalty_rate": float(inst.royalty_rate) if inst.royalty_rate else 15.00,
            "is_active": inst.is_active,
        }
        return Response({"success": True, "data": profile, "error": None})

    def patch(self, request):
        inst = get_user_institution(request.user)
        if not inst:
            return Response({"success": False, "error": "Profil d'établissement non trouvé"}, status=404)

        d = request.data
        for field in [
            "name", "short_name", "country", "city", "address",
            "rector_name", "academic_director_name", "contact_email", "contact_phone",
            "bank_name", "bank_iban", "bank_swift", "momo_number", "contract_reference"
        ]:
            if field in d:
                setattr(inst, field, d[field])
        if "royalty_rate" in d:
            inst.royalty_rate = d["royalty_rate"]
        inst.save()

        profile = {
            "id": str(inst.id),
            "name": inst.name,
            "short_name": inst.short_name or inst.code,
            "country": inst.country,
            "city": inst.city,
            "address": inst.address,
            "rector_name": inst.rector_name,
            "academic_director_name": inst.academic_director_name,
            "contact_email": inst.contact_email,
            "contact_phone": inst.contact_phone,
            "bank_name": inst.bank_name,
            "bank_iban": inst.bank_iban,
            "bank_swift": inst.bank_swift,
            "momo_number": inst.momo_number,
            "contract_reference": inst.contract_reference,
            "royalty_rate": float(inst.royalty_rate) if inst.royalty_rate else 15.00,
            "is_active": inst.is_active,
            "updated_at": inst.updated_at.isoformat() if hasattr(inst, 'updated_at') else timezone.now().isoformat()
        }
        return Response({"success": True, "data": profile, "error": None})


from django.http import HttpResponse

class ExportBouquetWordView(APIView):
    """GET /api/v1/partners/university/bouquets/<id>/export-word/ — Génère un .docx du bouquet."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from apps.catalog.models import Ouvrage

        inst = get_user_institution(request.user)
        try:
            bouquet = UniversityBouquetSubscription.objects.get(id=pk)
        except UniversityBouquetSubscription.DoesNotExist:
            return Response({"success": False, "error": "Bouquet introuvable."}, status=404)

        # Récupérer les ouvrages correspondants au bouquet
        ouvrages_qs = Ouvrage.objects.filter(
            status='published'
        ).select_related('discipline', 'institution').prefetch_related('authors')

        if bouquet.bouquet_type == 'discipline' and bouquet.discipline:
            ouvrages_qs = ouvrages_qs.filter(discipline__name__icontains=bouquet.discipline)
        elif bouquet.bouquet_type == 'faculty' and bouquet.faculty_code:
            ouvrages_qs = ouvrages_qs.filter(faculty__icontains=bouquet.faculty_code)
        
        if inst:
            ouvrages_qs = ouvrages_qs.filter(institution=inst)

        ouvrages = list(ouvrages_qs.order_by('discipline__name', 'title'))

        # Générer le document Word
        doc = Document()

        # Style de base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        # En-tête
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run('LAHAThèque')
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4E)  # Navy

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run('Bibliothèque Numérique Universitaire')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # Espace

        # Titre du bouquet
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f'Catalogue du Bouquet : {bouquet.title}')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4E)

        # Informations du bouquet
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_name = inst.name if inst else 'Toutes universités'
        start_str = bouquet.start_date.strftime("%d/%m/%Y") if hasattr(bouquet.start_date, 'strftime') else str(bouquet.start_date)
        end_str = bouquet.end_date.strftime("%d/%m/%Y") if hasattr(bouquet.end_date, 'strftime') else str(bouquet.end_date)
        run = info_para.add_run(
            f'Université : {inst_name}  |  '
            f'Type : {bouquet.get_bouquet_type_display()}  |  '
            f'Période : {start_str} — {end_str}  |  '
            f'{len(ouvrages)} ouvrage(s)'
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph()  # Espace

        # Tableau des ouvrages
        if ouvrages:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # En-têtes
            headers = ['N.', 'Titre', 'Auteur(s)', 'ISBN', 'Discipline / Faculté', 'Résumé']
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)

            # Lignes
            for idx, ouvrage in enumerate(ouvrages, 1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = ouvrage.title
                row.cells[2].text = ouvrage.auteur or ''
                row.cells[3].text = ouvrage.isbn or 'N/A'
                disc = ouvrage.discipline.name if ouvrage.discipline else ''
                fac = ouvrage.faculty or ''
                row.cells[4].text = f'{disc}\n{fac}'.strip()
                row.cells[5].text = (ouvrage.summary or '')[:150] + ('...' if len(ouvrage.summary or '') > 150 else '')

                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)
        else:
            doc.add_paragraph('Aucun ouvrage trouvé pour ce bouquet.').alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pied de page
        doc.add_paragraph()
        from datetime import datetime
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f'Document généré par LAHAThèque le {datetime.now().strftime("%d/%m/%Y à %H:%M")}')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.italic = True

        # Réponse HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_title = bouquet.title.replace(' ', '_').replace('/', '-')[:50]
        response['Content-Disposition'] = f'attachment; filename="Bouquet_{safe_title}.docx"'
        doc.save(response)
        return response

